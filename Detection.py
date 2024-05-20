# Import necessary libraries
import os
import shutil
import pandas as pd
import csv
import json
import cv2
import numpy as np
from sklearn.model_selection import GroupShuffleSplit
from detectron2.engine import DefaultTrainer, DefaultPredictor
from detectron2.config import get_cfg
from detectron2 import model_zoo
from detectron2.data.datasets import register_coco_instances
from detectron2.utils.logger import setup_logger
from detectron2.utils.visualizer import Visualizer, ColorMode
from detectron2.data import MetadataCatalog, DatasetCatalog
import matplotlib.pyplot as plt
import random
import torch
import copy
import albumentations as A
from albumentations.pytorch import ToTensorV2
from detectron2.data import detection_utils as utils
from detectron2.data import DatasetMapper, build_detection_train_loader
from docx import Document

# Set the sharing strategy to 'file_system'
torch.multiprocessing.set_sharing_strategy('file_system')

# Function to draw annotations on images and get image dimensions
def draw_annotations_and_get_dimensions(image_path, annotations, classification):
    image = cv2.imread(image_path)
    height, width = image.shape[:2]
    for annotation in annotations:
        points = list(zip(annotation['all_points_x'], annotation['all_points_y']))
        points = np.array(points, np.int32)
        points = points.reshape((-1, 1, 2))
        cv2.polylines(image, [points], isClosed=True, color=(255, 0, 0), thickness=2)
    
    cv2.putText(image, classification, (10, height - 10), cv2.FONT_HERSHEY_SIMPLEX, 1, (255, 0, 0), 2, cv2.LINE_AA)
    return image, width, height

# Load the annotations from the CSV file
segmentation_file = './data/Radiology_hand_drawn_segmentations_v2.csv'
annotations = []

with open(segmentation_file, newline='') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        if row['region_shape_attributes']:  # Only process rows with region_shape_attributes
            region_shape_attributes = json.loads(row['region_shape_attributes'])
            annotations.append({
                'filename': row['#filename'],
                'points': {
                    'all_points_x': region_shape_attributes.get('all_points_x', []),
                    'all_points_y': region_shape_attributes.get('all_points_y', [])
                }
            })

# Group annotations by filename
annotations_by_filename = {}
for annotation in annotations:
    if annotation['filename'] not in annotations_by_filename:
        annotations_by_filename[annotation['filename']] = []
    annotations_by_filename[annotation['filename']].append(annotation['points'])

# Get the list of annotated images
annotated_images = list(annotations_by_filename.keys())

# Load the Excel file for classifications
annotations_file = './data/Radiology-manual-annotations.xlsx'
df_annotations = pd.read_excel(annotations_file)

# Classify images based on the annotations
def classify_images(df):
    classifications = {}
    for index, row in df.iterrows():
        image_name = row['Image_name']
        classification = row['Pathology Classification/ Follow up']
        
        if classification == 'Benign':
            classifications[image_name] = 'Benign'
        elif classification == 'Malignant':
            classifications[image_name] = 'Malignant'
        else:
            classifications[image_name] = 'Normal'
    
    return classifications

# Get the classifications
classifications = classify_images(df_annotations)

# Ensure directories exist
os.makedirs('./data/annotated_images', exist_ok=True)

# Annotate images, save them, and collect their dimensions
images_info = []
annotations_info = []
category_id = 1  # Assuming one category for simplicity

for filename, points in annotations_by_filename.items():
    image_path = f'./data/images/{filename}'
    if os.path.exists(image_path):
        classification = classifications.get(os.path.splitext(filename)[0], 'Normal')
        image_with_annotations, width, height = draw_annotations_and_get_dimensions(image_path, points, classification)
        annotated_image_path = f'./data/annotated_images/{filename}'
        cv2.imwrite(annotated_image_path, image_with_annotations)
        
        image_id = len(images_info) + 1
        images_info.append({
            "id": image_id,
            "file_name": filename,
            "width": int(width),  # Convert to int
            "height": int(height)  # Convert to int
        })
        
        for annotation in points:
            all_points_x = annotation['all_points_x']
            all_points_y = annotation['all_points_y']
            if all_points_x and all_points_y:  # Ensure the points are not empty
                segmentation = [list(np.array([all_points_x, all_points_y]).T.flatten())]
                bbox = [int(min(all_points_x)), int(min(all_points_y)), int(max(all_points_x) - min(all_points_x)), int(max(all_points_y) - min(all_points_y))]  # Convert to int
                area = int(cv2.contourArea(np.array(list(zip(all_points_x, all_points_y)), dtype=np.int32)))  # Convert to int
                
                annotations_info.append({
                    "id": len(annotations_info) + 1,
                    "image_id": image_id,
                    "category_id": category_id,
                    "segmentation": segmentation,
                    "area": area,
                    "bbox": bbox,
                    "iscrowd": 0
                })

categories_info = [{
    "id": category_id,
    "name": "Lesion"  # Change this to the actual category name
}]

# Get all images including normal ones
all_images = [f for f in os.listdir('./data/images') if f.endswith(('.jpg', '.jpeg'))]

# Include normal images
for image_file in all_images:
    if image_file not in annotated_images:
        image_path = f'./data/images/{image_file}'
        if os.path.exists(image_path):
            image = cv2.imread(image_path)
            height, width = image.shape[:2]
            
            image_id = len(images_info) + 1
            images_info.append({
                "id": image_id,
                "file_name": image_file,
                "width": int(width),  # Convert to int
                "height": int(height)  # Convert to int
            })

# Create the COCO JSON annotation file
coco_annotation = {
    "images": images_info,
    "annotations": annotations_info,
    "categories": categories_info
}

# Ensure all numpy types are converted to native Python types
def convert_to_native_types(data):
    if isinstance(data, list):
        return [convert_to_native_types(item) for item in data]
    elif isinstance(data, dict):
        return {key: convert_to_native_types(value) for key, value in data.items()}
    elif isinstance(data, np.generic):
        return data.item()
    else:
        return data

coco_annotation = convert_to_native_types(coco_annotation)

with open('./data/annotated_images/annotations.json', 'w') as f:
    json.dump(coco_annotation, f)

# Split the images into train and test sets
patient_groups = {}
for image_info in images_info:
    patient_id = image_info['file_name'].split('_')[0]
    if patient_id not in patient_groups:
        patient_groups[patient_id] = []
    patient_groups[patient_id].append(image_info)

group_kfold = GroupShuffleSplit(n_splits=1, test_size=0.2, random_state=42)
patient_ids = list(patient_groups.keys())
train_idx, test_idx = next(group_kfold.split(X=patient_ids, groups=patient_ids))

train_patient_ids = [patient_ids[idx] for idx in train_idx]
test_patient_ids = [patient_ids[idx] for idx in test_idx]

train_images = [img for pid in train_patient_ids for img in patient_groups[pid]]
test_images = [img for pid in test_patient_ids for img in patient_groups[pid]]

# Create new COCO JSON files for train and test sets
def create_coco_json(images, annotations, categories, dest_file):
    image_ids = [img['id'] for img in images]
    filtered_annotations = [anno for anno in annotations if anno['image_id'] in image_ids]
    coco_data = {
        "images": images,
        "annotations": filtered_annotations,
        "categories": categories
    }
    coco_data = convert_to_native_types(coco_data)  # Convert to native types before saving
    with open(dest_file, 'w') as f:
        json.dump(coco_data, f)

# Ensure the destination directories exist
os.makedirs('./data/train', exist_ok=True)
os.makedirs('./data/valid', exist_ok=True)

# Create train and test COCO JSON files
create_coco_json(train_images, annotations_info, categories_info, './data/train/_annotations.coco.json')
create_coco_json(test_images, annotations_info, categories_info, './data/valid/_annotations.coco.json')

# Move the train and test images to their respective directories
def move_files(images, src_folder, dest_folder):
    os.makedirs(dest_folder, exist_ok=True)
    for img in images:
        src = os.path.join(src_folder, img['file_name'])
        dest = os.path.join(dest_folder, img['file_name'])
        shutil.copyfile(src, dest)

move_files(train_images, './data/images', './data/train')
move_files(test_images, './data/images', './data/valid')

# Count the number of each type in train and test sets
train_counts = {'Benign': 0, 'Malignant': 0, 'Normal': 0}
test_counts = {'Benign': 0, 'Malignant': 0, 'Normal': 0}

for image in train_images:
    image_name = os.path.splitext(image['file_name'])[0]  # Get the image name without extension
    classification = classifications.get(image_name, 'Normal')
    train_counts[classification] += 1

for image in test_images:
    image_name = os.path.splitext(image['file_name'])[0]  # Get the image name without extension
    classification = classifications.get(image_name, 'Normal')
    test_counts[classification] += 1

print("Train counts:", train_counts)
print("Test counts:", test_counts)

# Register datasets
register_coco_instances("my_dataset_train", {}, "./data/train/_annotations.coco.json", "./data/train")
register_coco_instances("my_dataset_val", {}, "./data/valid/_annotations.coco.json", "./data/valid")

# Set up the logger
setup_logger()

# Define Albumentations augmentations
def get_albumentations_transforms():
    return A.Compose([
        A.Rotate(limit=45, p=0.5),  # Rotate the image by up to 45 degrees
        A.HorizontalFlip(p=0.5),    # Horizontally flip the image
        A.VerticalFlip(p=0.5),      # Vertically flip the image
        A.RandomBrightnessContrast(p=0.5),  # Randomly change the brightness and contrast
        A.ShiftScaleRotate(shift_limit=0.0625, scale_limit=0.2, rotate_limit=45, p=0.5),
        A.OneOf([
            A.HueSaturationValue(10,15,10),
            A.RandomBrightnessContrast(brightness_limit=0.3, contrast_limit=0.3)
        ], p=0.5),
        A.MotionBlur(blur_limit=3, p=0.2),
        A.IAAPerspective(scale=(0.02, 0.05), p=0.3),
        A.Cutout(num_holes=8, max_h_size=8, max_w_size=8, p=0.3),
        A.Normalize(mean=(0.485, 0.456, 0.406), std=(0.229, 0.224, 0.225)),
        ToTensorV2()
    ])

# Define Albumentations Mapper
class AlbumentationsMapper(DatasetMapper):
    def __init__(self, cfg, is_train=True):
        super().__init__(cfg, is_train)
        self.augmentations = get_albumentations_transforms()

    def __call__(self, dataset_dict):
        dataset_dict = copy.deepcopy(dataset_dict)
        image = utils.read_image(dataset_dict["file_name"], format=self.img_format)
        aug_output = self.augmentations(image=image)
        image = aug_output["image"]

        dataset_dict["image"] = torch.as_tensor(image.transpose(2, 0, 1).astype("float32"))

        annos = [utils.transform_instance_annotations(annotation, [], image.shape[:2])
                 for annotation in dataset_dict.pop("annotations")]
        instances = utils.annotations_to_instances(annos, image.shape[:2])
        dataset_dict["instances"] = utils.filter_empty_instances(instances)
        return dataset_dict

# Define Custom Trainer with Albumentations
class TrainerWithCustomLoader(DefaultTrainer):
    @classmethod
    def build_train_loader(cls, cfg):
        return build_detection_train_loader(cfg, mapper=AlbumentationsMapper(cfg, is_train=True))

# Configuration setup
cfg = get_cfg()
cfg.merge_from_file(model_zoo.get_config_file("COCO-Detection/faster_rcnn_R_50_FPN_3x.yaml"))
cfg.DATASETS.TRAIN = ("my_dataset_train",)
cfg.DATASETS.TEST = ("my_dataset_val",)
cfg.DATALOADER.NUM_WORKERS = 0  # Disable multiprocessing to avoid shared memory issue
cfg.MODEL.WEIGHTS = model_zoo.get_checkpoint_url("COCO-Detection/faster_rcnn_R_50_FPN_3x.yaml")
cfg.SOLVER.IMS_PER_BATCH = 4  # Increase batch size
cfg.SOLVER.BASE_LR = 0.0025  # Increase learning rate
cfg.SOLVER.MAX_ITER = 3000  # Increase max iterations
cfg.SOLVER.STEPS = []  # do not decay learning rate
cfg.MODEL.ROI_HEADS.BATCH_SIZE_PER_IMAGE = 256  # Increase batch size per image
cfg.MODEL.ROI_HEADS.NUM_CLASSES = 3  # Update this based on your classes (Benign, Malignant, Normal)
cfg.MODEL.DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
cfg.TEST.EVAL_PERIOD = 500  # Evaluate every 500 iterations

os.makedirs(cfg.OUTPUT_DIR, exist_ok=True)
trainer = TrainerWithCustomLoader(cfg)
trainer.resume_or_load(resume=False)
trainer.train()

# Perform inference with the trained model
cfg.MODEL.WEIGHTS = os.path.join(cfg.OUTPUT_DIR, "model_final.pth")
cfg.MODEL.ROI_HEADS.SCORE_THRESH_TEST = 0.5
predictor = DefaultPredictor(cfg)

# Load the dataset
dataset_dicts = DatasetCatalog.get("my_dataset_val")
metadata = MetadataCatalog.get("my_dataset_val")

# Function to calculate metrics and find the best threshold
def calculate_metrics_and_find_best_threshold(predictor, dataset_dicts, metadata):
    thresholds = np.arange(0.1, 1.0, 0.1)
    best_threshold = 0.5
    best_f1 = 0.0

    results = []

    for threshold in thresholds:
        cfg.MODEL.ROI_HEADS.SCORE_THRESH_TEST = float(threshold)
        predictor = DefaultPredictor(cfg)

        y_true = []
        y_pred = []

        for d in dataset_dicts:
            im = cv2.imread(d["file_name"])
            outputs = predictor(im)
            instances = outputs["instances"].to("cpu")
            classes = instances.pred_classes.tolist()
            scores = instances.scores.tolist()

            for class_id, score in zip(classes, scores):
                y_pred.append(class_id)
                y_true.append(d['annotations'][0]['category_id'])

        f1 = f1_score(y_true, y_pred, average='weighted')
        accuracy = accuracy_score(y_true, y_pred)

        results.append({
            'Threshold': float(threshold),
            'F1 Score': f1,
            'Accuracy': accuracy
        })

        if f1 > best_f1:
            best_f1 = f1
            best_threshold = float(threshold)

    return best_threshold, results

# Calculate metrics and find the best threshold
best_threshold, results = calculate_metrics_and_find_best_threshold(predictor, dataset_dicts, metadata)
cfg.MODEL.ROI_HEADS.SCORE_THRESH_TEST = best_threshold
predictor = DefaultPredictor(cfg)

print(f"Best threshold: {best_threshold}")

# Create a Word document to store results
doc = Document()
doc.add_heading('Inference Results', level=1)

# Display results for 5 random images and save to the document
results_table = []
for d in random.sample(dataset_dicts, 5):
    im = cv2.imread(d["file_name"])
    outputs = predictor(im)
    instances = outputs["instances"].to("cpu")
    classes = instances.pred_classes.tolist()
    scores = instances.scores.tolist()

    # Store predictions and ground truth
    ground_truth_class = d['annotations'][0]['category_id']
    predicted_class = classes[0] if classes else 'None'
    predicted_score = scores[0] if scores else 'None'
    results_table.append([d["file_name"], ground_truth_class, predicted_class, predicted_score])

    # Modify Visualizer to display classification
    v = Visualizer(im[:, :, ::-1], metadata=metadata, scale=0.8)
    out = v.draw_instance_predictions(outputs["instances"].to("cpu"))
    
    # Save the image with predictions
    result_image_path = f'./data/results/{os.path.basename(d["file_name"])}'
    os.makedirs(os.path.dirname(result_image_path), exist_ok=True)
    cv2.imwrite(result_image_path, out.get_image()[:, :, ::-1])
    
    # Add the result to the document
    doc.add_heading(f'Image: {d["file_name"]}', level=2)
    doc.add_paragraph(f'Ground Truth: {ground_truth_class}')
    doc.add_paragraph(f'Predicted Class: {predicted_class}')
    doc.add_paragraph(f'Predicted Score: {predicted_score}')
    doc.add_picture(result_image_path, width=docx.shared.Inches(4))

# Save the document
doc.save('./data/results/inference_results.docx')

# Convert results to DataFrame and display
results_df = pd.DataFrame(results_table, columns=["File Name", "Ground Truth", "Predicted Class", "Predicted Score"])
print(results_df)

# Display the evaluation results
eval_df = pd.DataFrame(results)
print(eval_df)

# Calculate the overall F1 score and accuracy with the best threshold
y_true = []
y_pred = []
for d in dataset_dicts:
    im = cv2.imread(d["file_name"])
    outputs = predictor(im)
    instances = outputs["instances"].to("cpu")
    classes = instances.pred_classes.tolist()

    for class_id in classes:
        y_pred.append(class_id)
        y_true.append(d['annotations'][0]['category_id'])

final_f1 = f1_score(y_true, y_pred, average='weighted')
final_accuracy = accuracy_score(y_true, y_pred)
print(f"Final F1 Score: {final_f1}")
print(f"Final Accuracy: {final_accuracy}")
